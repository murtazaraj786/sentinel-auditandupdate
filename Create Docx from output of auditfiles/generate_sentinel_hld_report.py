
import argparse
import os
import pandas as pd
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def shade_cell(cell, shade_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), shade_hex)  # e.g., "D9D9D9"
    tcPr.append(shd)

def add_styled_table(doc, df, max_rows=10):
    """Add a styled table with header shading and alternating row stripes."""
    if df.empty:
        doc.add_paragraph("No data found.")
        return

    preview = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(preview.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(preview.columns):
        hdr_cells[i].text = str(col)
        p = hdr_cells[i].paragraphs[0]
        if p.runs:
            p.runs[0].font.bold = True
        else:
            run = p.add_run()
            run.bold = True
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(hdr_cells[i], "D9D9D9")  # header grey

    # Rows
    for idx, row in preview.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = "" if pd.isna(val) else str(val)
            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # Alternate row shading
        if idx % 2 == 0:
            for cell in row_cells:
                shade_cell(cell, "F2F2F2")

def add_section(doc, title, df, notes=None):
    doc.add_heading(title, level=2)
    doc.add_paragraph(f"Total Records: {len(df)}")
    if notes:
        doc.add_paragraph(notes)
    doc.add_paragraph("Preview of Key Data:")
    add_styled_table(doc, df)

def infer_enabled_count(df):
    for col in df.columns:
        if col.strip().lower() in ("enabled", "is_enabled", "ruleenabled"):
            try:
                return int(df[col].astype(str).str.lower().isin(["true","1","yes","y"]).sum())
            except Exception:
                try:
                    return int(df[col].sum())
                except Exception:
                    return None
    return None

def detect_customer_name_from_metadata():
    """Try to detect customer name from metadata files in parent directories."""
    import glob
    
    # Look for metadata files in common locations
    search_paths = [
        "../Sentinel Audit/sentinel_customer_info_*.csv",
        "../Sentinel SOC Optimisation Audit/soc_customer_info_*.csv",
        "./sentinel_customer_info_*.csv",
        "./soc_customer_info_*.csv"
    ]
    
    for pattern in search_paths:
        metadata_files = glob.glob(pattern)
        if metadata_files:
            # Use the most recent metadata file
            latest_file = max(metadata_files, key=os.path.getmtime)
            try:
                metadata_df = pd.read_csv(latest_file)
                if not metadata_df.empty and 'customer_name' in metadata_df.columns:
                    customer_name = metadata_df.iloc[0]['customer_name']
                    print(f"🔍 Auto-detected customer name from {latest_file}: {customer_name}")
                    return customer_name
            except Exception as e:
                print(f"⚠️  Could not read metadata from {latest_file}: {e}")
    
    return None

def main():
    parser = argparse.ArgumentParser(description="Generate a Microsoft Sentinel HLD-style audit report from CSV exports.")
    parser.add_argument("--analytic-rules", required=True, help="Path to sentinel analytic rules CSV")
    parser.add_argument("--data-connectors", required=True, help="Path to sentinel data connectors CSV")
    parser.add_argument("--soc-ingestion", required=True, help="Path to SOC data ingestion CSV")
    parser.add_argument("--soc-recommendations", required=True, help="Path to SOC recommendations CSV")
    parser.add_argument("--soc-rule-efficiency", required=True, help="Path to SOC rule efficiency CSV")
    parser.add_argument("--customer-name", help="Customer name to place in the report header (auto-detected if not specified)")
    parser.add_argument("--output", default="sentinel_hld_audit.docx", help="Output .docx path")
    parser.add_argument("--preview-rows", type=int, default=10, help="Rows to preview in each table")
    args = parser.parse_args()
    
    # Auto-detect customer name if not provided
    if not args.customer_name:
        detected_name = detect_customer_name_from_metadata()
        args.customer_name = detected_name if detected_name else "Azure Customer"

    # Load CSVs
    dfs = {
        "Analytic Rules": pd.read_csv(args.analytic_rules),
        "Data Connectors": pd.read_csv(args.data_connectors),
        "SOC Data Ingestion": pd.read_csv(args.soc_ingestion),
        "SOC Recommendations": pd.read_csv(args.soc_recommendations),
        "SOC Rule Efficiency": pd.read_csv(args.soc_rule_efficiency),
    }

    # Create document
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri' # type: ignore
    style.font.size = Pt(11) # type: ignore

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    doc.add_heading(f"{args.customer_name} – Microsoft Sentinel High-Level Design (HLD) Audit Report", level=1)
    doc.add_paragraph(f"Generated: {now}")
    doc.add_paragraph(
        "This HLD report assesses the Microsoft Sentinel deployment, focusing on data connector health, "
        "analytic rule coverage, SOC recommendations, and operational efficiency. It includes risk ratings and "
        "prioritised remediation actions based on supplied exports."
    )

    # Executive summary
    doc.add_heading("1. Executive Summary", level=2)
    ar = dfs["Analytic Rules"]
    dc = dfs["Data Connectors"]
    sr = dfs["SOC Recommendations"]
    reff = dfs["SOC Rule Efficiency"]
    di = dfs["SOC Data Ingestion"]

    enabled_count = infer_enabled_count(ar)
    doc.add_paragraph(
        f"Sentinel is operational with {len(dc)} data connectors and {len(ar)} analytic rules"
        + (f", {enabled_count} enabled." if enabled_count is not None else ".")
    )
    doc.add_paragraph(
        "Overall risk: Medium — improvements recommended in connector completeness, rule enablement consistency, "
        "and ingestion coverage."
    )

    # Data Connectors
    doc.add_heading("2. Data Connectors", level=2)
    doc.add_paragraph("Risk: Medium – Validate core sources (Identity, Email, Endpoint, Perimeter) are healthy and sending.")
    add_section(doc, "2.1 Connector Inventory", dc, notes=None)

    # Analytic Rules
    doc.add_heading("3. Analytic Rules", level=2)
    doc.add_paragraph("Risk: Medium – Confirm coverage vs. MITRE ATT&CK and ensure priority rules are enabled.")
    add_section(doc, "3.1 Analytic Rules Overview", ar, notes=None)

    # SOC Recommendations
    doc.add_heading("4. SOC Recommendations", level=2)
    doc.add_paragraph("Risk: Medium – Action outstanding recommendations and track to closure.")
    add_section(doc, "4.1 Recommendation Items", sr, notes=None)

    # SOC Rule Efficiency
    doc.add_heading("5. SOC Rule Efficiency", level=2)
    doc.add_paragraph("Risk: Low – Continue tuning noisy rules and enriching detections.")
    add_section(doc, "5.1 Efficiency Metrics", reff, notes=None)

    # SOC Data Ingestion
    doc.add_heading("6. SOC Data Ingestion", level=2)
    doc.add_paragraph("Risk: Medium – Review gaps in ingestion and ensure critical tables are present and up to date.")
    add_section(doc, "6.1 Ingestion Overview", di, notes=None)

    # Save
    out_path = args.output
    doc.save(out_path)
    print(out_path)

if __name__ == "__main__":
    main()
