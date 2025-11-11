
import argparse
import os
import pandas as pd
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def shade_cell(cell, shade_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), shade_hex)  # e.g., "D9D9D9"
    tcPr.append(shd)

def add_styled_table(doc, df, max_rows=10):
    """Add a styled table with header shading and alternating row stripes."""
    if df.empty:
        doc.add_paragraph("No data found.")
        return

    preview = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(preview.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(preview.columns):
        hdr_cells[i].text = str(col)
        p = hdr_cells[i].paragraphs[0]
        if p.runs:
            p.runs[0].font.bold = True
        else:
            run = p.add_run()
            run.bold = True
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(hdr_cells[i], "D9D9D9")  # header grey

    # Rows
    for idx, row in preview.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = "" if pd.isna(val) else str(val)
            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # Alternate row shading
        if idx % 2 == 0:
            for cell in row_cells:
                shade_cell(cell, "F2F2F2")

def add_section(doc, title, df, notes=None):
    doc.add_heading(title, level=2)
    doc.add_paragraph(f"Total Records: {len(df)}")
    if notes:
        doc.add_paragraph(notes)
    doc.add_paragraph("Preview of Key Data:")
    add_styled_table(doc, df)

def infer_enabled_count(df):
    for col in df.columns:
        if col.strip().lower() in ("enabled", "is_enabled", "ruleenabled"):
            try:
                return int(df[col].astype(str).str.lower().isin(["true","1","yes","y"]).sum())
            except Exception:
                try:
                    return int(df[col].sum())
                except Exception:
                    return None
    return None

def detect_customer_name_from_metadata():
    """Try to detect customer name from metadata files in parent directories."""
    import glob
    
    # Look for metadata files in common locations
    search_paths = [
        "../Sentinel Audit/sentinel_customer_info_*.csv",
        "../Sentinel SOC Optimisation Audit/soc_customer_info_*.csv",
        "../Defender XDR Audit/defender_xdr_*.csv",
        "./sentinel_customer_info_*.csv",
        "./soc_customer_info_*.csv",
        "./defender_xdr_*.csv"
    ]
    
    for pattern in search_paths:
        metadata_files = glob.glob(pattern)
        if metadata_files:
            # Use the most recent metadata file
            latest_file = max(metadata_files, key=os.path.getmtime)
            try:
                metadata_df = pd.read_csv(latest_file)
                if not metadata_df.empty and 'customer_name' in metadata_df.columns:
                    customer_name = str(metadata_df.iloc[0]['customer_name']).strip()
                    if customer_name:
                        print(f"🔍 Auto-detected customer name from {latest_file}: {customer_name}")
                        return customer_name
            except Exception as e:
                print(f"⚠️  Could not read metadata from {latest_file}: {e}")
    
    return None

def auto_detect_xdr_files():
    """Auto-detect Defender XDR CSV files in current and parent directories."""
    import glob
    
    detected_files = {}
    
    search_patterns = {
        "xdr_security_alerts": [
            "../Defender XDR Audit/defender_xdr_security_alerts_*.csv",
            "./defender_xdr_security_alerts_*.csv"
        ],
        "xdr_security_incidents": [
            "../Defender XDR Audit/defender_xdr_security_incidents_*.csv", 
            "./defender_xdr_security_incidents_*.csv"
        ],
        "xdr_attack_simulations": [
            "../Defender XDR Audit/defender_xdr_attack_simulations_*.csv",
            "./defender_xdr_attack_simulations_*.csv"
        ],
        "xdr_secure_score": [
            "../Defender XDR Audit/defender_xdr_secure_score_*.csv",
            "./defender_xdr_secure_score_*.csv"
        ]
    }
    
    for file_type, patterns in search_patterns.items():
        for pattern in patterns:
            files = glob.glob(pattern)
            if files:
                # Use the most recent file
                latest_file = max(files, key=os.path.getmtime)
                detected_files[file_type] = latest_file
                print(f"🔍 Auto-detected {file_type}: {latest_file}")
                break
    
    return detected_files

def main():
    parser = argparse.ArgumentParser(description="Generate a comprehensive Microsoft Sentinel and Defender XDR HLD-style audit report from CSV exports.")
    
    # Sentinel required arguments
    parser.add_argument("--analytic-rules", required=True, help="Path to sentinel analytic rules CSV")
    parser.add_argument("--data-connectors", required=True, help="Path to sentinel data connectors CSV")
    parser.add_argument("--soc-ingestion", required=True, help="Path to SOC data ingestion CSV")
    parser.add_argument("--soc-recommendations", required=True, help="Path to SOC recommendations CSV")
    parser.add_argument("--soc-rule-efficiency", required=True, help="Path to SOC rule efficiency CSV")
    
    # Defender XDR optional arguments
    parser.add_argument("--xdr-security-alerts", help="Path to Defender XDR security alerts CSV")
    parser.add_argument("--xdr-security-incidents", help="Path to Defender XDR security incidents CSV")
    parser.add_argument("--xdr-attack-simulations", help="Path to Defender XDR attack simulations CSV")
    parser.add_argument("--xdr-secure-score", help="Path to Defender XDR secure score CSV")
    
    # General arguments
    parser.add_argument("--customer-name", help="Customer name to place in the report header (auto-detected if not specified)")
    parser.add_argument("--output", default="sentinel_xdr_hld_audit.docx", help="Output .docx path")
    parser.add_argument("--preview-rows", type=int, default=10, help="Rows to preview in each table")
    args = parser.parse_args()
    
    # Auto-detect customer name if not provided
    if not args.customer_name:
        detected_name = detect_customer_name_from_metadata()
        args.customer_name = detected_name if detected_name else "Azure Customer"
    
    # Auto-detect XDR files if not explicitly provided
    auto_detected_xdr = auto_detect_xdr_files()
    if not args.xdr_security_alerts and "xdr_security_alerts" in auto_detected_xdr:
        args.xdr_security_alerts = auto_detected_xdr["xdr_security_alerts"]
    if not args.xdr_security_incidents and "xdr_security_incidents" in auto_detected_xdr:
        args.xdr_security_incidents = auto_detected_xdr["xdr_security_incidents"]
    if not args.xdr_attack_simulations and "xdr_attack_simulations" in auto_detected_xdr:
        args.xdr_attack_simulations = auto_detected_xdr["xdr_attack_simulations"]
    if not args.xdr_secure_score and "xdr_secure_score" in auto_detected_xdr:
        args.xdr_secure_score = auto_detected_xdr["xdr_secure_score"]

    # Load Sentinel CSVs
    dfs = {
        "Analytic Rules": pd.read_csv(args.analytic_rules),
        "Data Connectors": pd.read_csv(args.data_connectors),
        "SOC Data Ingestion": pd.read_csv(args.soc_ingestion),
        "SOC Recommendations": pd.read_csv(args.soc_recommendations),
        "SOC Rule Efficiency": pd.read_csv(args.soc_rule_efficiency),
    }
    
    # Load optional Defender XDR CSVs
    xdr_dfs = {}
    if args.xdr_security_alerts:
        try:
            xdr_dfs["XDR Security Alerts"] = pd.read_csv(args.xdr_security_alerts)
        except Exception as e:
            print(f"⚠️  Could not load XDR Security Alerts: {e}")
    
    if args.xdr_security_incidents:
        try:
            xdr_dfs["XDR Security Incidents"] = pd.read_csv(args.xdr_security_incidents)
        except Exception as e:
            print(f"⚠️  Could not load XDR Security Incidents: {e}")
    
    if args.xdr_attack_simulations:
        try:
            xdr_dfs["XDR Attack Simulations"] = pd.read_csv(args.xdr_attack_simulations)
        except Exception as e:
            print(f"⚠️  Could not load XDR Attack Simulations: {e}")
    
    if args.xdr_secure_score:
        try:
            xdr_dfs["XDR Secure Score"] = pd.read_csv(args.xdr_secure_score)
        except Exception as e:
            print(f"⚠️  Could not load XDR Secure Score: {e}")
    
    # Check if we have any XDR data
    include_xdr = len(xdr_dfs) > 0

    # Create document
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri' # type: ignore
    style.font.size = Pt(11) # type: ignore

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    title_suffix = " and Defender XDR" if include_xdr else ""
    doc.add_heading(f"{args.customer_name} – Microsoft Sentinel{title_suffix} High-Level Design (HLD) Audit Report", level=1)
    doc.add_paragraph(f"Generated: {now}")
    
    description = (
        "This HLD report assesses the Microsoft Sentinel deployment, focusing on data connector health, "
        "analytic rule coverage, SOC recommendations, and operational efficiency."
    )
    if include_xdr:
        description += " It also includes Microsoft Defender XDR security posture, including alerts, incidents, attack simulations, and secure score metrics."
    description += " The report includes risk ratings and prioritised remediation actions based on supplied exports."
    
    doc.add_paragraph(description)

    # Executive summary
    doc.add_heading("1. Executive Summary", level=2)
    ar = dfs["Analytic Rules"]
    dc = dfs["Data Connectors"]
    sr = dfs["SOC Recommendations"]
    reff = dfs["SOC Rule Efficiency"]
    di = dfs["SOC Data Ingestion"]

    enabled_count = infer_enabled_count(ar)
    summary_text = (
        f"Microsoft Sentinel is operational with {len(dc)} data connectors and {len(ar)} analytic rules"
        + (f", {enabled_count} enabled." if enabled_count is not None else ".")
    )
    
    if include_xdr:
        xdr_summary_parts = []
        if "XDR Security Alerts" in xdr_dfs:
            xdr_summary_parts.append(f"{len(xdr_dfs['XDR Security Alerts'])} security alerts")
        if "XDR Security Incidents" in xdr_dfs:
            xdr_summary_parts.append(f"{len(xdr_dfs['XDR Security Incidents'])} security incidents")
        if "XDR Attack Simulations" in xdr_dfs:
            xdr_summary_parts.append(f"{len(xdr_dfs['XDR Attack Simulations'])} attack simulations")
        
        if xdr_summary_parts:
            summary_text += f" Defender XDR shows {', '.join(xdr_summary_parts)}."
    
    doc.add_paragraph(summary_text)
    
    risk_assessment = (
        "Overall risk: Medium — improvements recommended in connector completeness, rule enablement consistency, "
        "and ingestion coverage."
    )
    if include_xdr:
        risk_assessment += " XDR integration provides enhanced threat detection and response capabilities."
    
    doc.add_paragraph(risk_assessment)

    # Data Connectors
    doc.add_heading("2. Data Connectors", level=2)
    doc.add_paragraph("Risk: Medium – Validate core sources (Identity, Email, Endpoint, Perimeter) are healthy and sending.")
    add_section(doc, "2.1 Connector Inventory", dc, notes=None)

    # Analytic Rules
    doc.add_heading("3. Analytic Rules", level=2)
    doc.add_paragraph("Risk: Medium – Confirm coverage vs. MITRE ATT&CK and ensure priority rules are enabled.")
    add_section(doc, "3.1 Analytic Rules Overview", ar, notes=None)

    # SOC Recommendations
    doc.add_heading("4. SOC Recommendations", level=2)
    doc.add_paragraph("Risk: Medium – Action outstanding recommendations and track to closure.")
    add_section(doc, "4.1 Recommendation Items", sr, notes=None)

    # SOC Rule Efficiency
    doc.add_heading("5. SOC Rule Efficiency", level=2)
    doc.add_paragraph("Risk: Low – Continue tuning noisy rules and enriching detections.")
    add_section(doc, "5.1 Efficiency Metrics", reff, notes=None)

    # SOC Data Ingestion
    doc.add_heading("6. SOC Data Ingestion", level=2)
    doc.add_paragraph("Risk: Medium – Review gaps in ingestion and ensure critical tables are present and up to date.")
    add_section(doc, "6.1 Ingestion Overview", di, notes=None)

    # Defender XDR Sections (if data is available)
    if include_xdr:
        section_num = 7
        
        if "XDR Security Alerts" in xdr_dfs:
            doc.add_heading(f"{section_num}. Defender XDR Security Alerts", level=2)
            alerts_df = xdr_dfs["XDR Security Alerts"]
            
            # Calculate severity breakdown if severity column exists
            severity_note = None
            if 'Severity' in alerts_df.columns:
                severity_counts = alerts_df['Severity'].value_counts()
                severity_note = f"Severity breakdown: {', '.join([f'{k}: {v}' for k, v in severity_counts.items()])}"
            
            doc.add_paragraph("Risk: Medium – Review and triage active security alerts, prioritizing high/critical severity items.")
            add_section(doc, f"{section_num}.1 Active Security Alerts", alerts_df, notes=severity_note)
            section_num += 1
        
        if "XDR Security Incidents" in xdr_dfs:
            doc.add_heading(f"{section_num}. Defender XDR Security Incidents", level=2)
            incidents_df = xdr_dfs["XDR Security Incidents"]
            
            # Calculate status breakdown if status column exists
            status_note = None
            if 'Status' in incidents_df.columns:
                status_counts = incidents_df['Status'].value_counts()
                status_note = f"Status breakdown: {', '.join([f'{k}: {v}' for k, v in status_counts.items()])}"
            
            doc.add_paragraph("Risk: Medium – Ensure incidents are properly investigated and resolved in timely manner.")
            add_section(doc, f"{section_num}.1 Security Incidents Overview", incidents_df, notes=status_note)
            section_num += 1
        
        if "XDR Attack Simulations" in xdr_dfs:
            doc.add_heading(f"{section_num}. Defender XDR Attack Simulations", level=2)
            simulations_df = xdr_dfs["XDR Attack Simulations"]
            doc.add_paragraph("Risk: Low – Continue phishing simulations to maintain user awareness and identify training gaps.")
            add_section(doc, f"{section_num}.1 Attack Simulation Campaigns", simulations_df, notes=None)
            section_num += 1
        
        if "XDR Secure Score" in xdr_dfs:
            doc.add_heading(f"{section_num}. Defender XDR Secure Score", level=2)
            score_df = xdr_dfs["XDR Secure Score"]
            
            # Extract current score if available
            score_note = None
            if 'Percentage' in score_df.columns and not score_df.empty:
                current_score = score_df.iloc[0]['Percentage']
                score_note = f"Current security score: {current_score}"
            
            doc.add_paragraph("Risk: Low-Medium – Continue improving security posture based on secure score recommendations.")
            add_section(doc, f"{section_num}.1 Security Posture Score", score_df, notes=score_note)

    # Save
    out_path = args.output
    doc.save(out_path)
    print(out_path)

if __name__ == "__main__":
    main()
